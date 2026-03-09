"""
sop_parser.py — Parse an SOP .docx and extract structured process information.

Pipeline
--------
1. Load .docx with python-docx (fallback: plain text via zipfile).
2. Split document into steps / sections using heading styles.
3. For each section, run:
   a. ISA tag regex extraction (FCV-101, PT-203 …)
   b. Valve position extraction ("open FCV-101", "close XV-200")
   c. Operating parameter extraction (pressure, temperature, flow)
   d. Required component enumeration
4. Produce list[SopStep] for downstream cross-referencing.

Optional NLP Enhancement
------------------------
If spaCy is installed, a custom NER pipeline can be layered on top of
the regex extractor for higher recall on complex sentences.
"""

from __future__ import annotations

import logging
import re
import zipfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from pid_graph.config import SopConfig
from pid_graph.models import SopStep

log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Regex patterns
# ---------------------------------------------------------------------------

# ISA tag:  FCV-101  PT_3A  LT-202B  P-101
_ISA_PAT = re.compile(r"\b([A-Z]{1,4})[_\-]?(\d{2,5}[A-Z]?)\b")

# Valve state commands
_OPEN_PAT   = re.compile(r"\b(?:open|opened?|opening)\b[^A-Z]{0,15}([A-Z]{1,4}[_\-]?\d{2,5}[A-Z]?)", re.I)
_CLOSE_PAT  = re.compile(r"\b(?:close|closed?|closing|shut|shutoff)\b[^A-Z]{0,15}([A-Z]{1,4}[_\-]?\d{2,5}[A-Z]?)", re.I)

# Operating parameters
_PRESSURE_PAT = re.compile(r"(\d+(?:\.\d+)?)\s*(psi[ga]?|bar[ag]?|kpa|mpa|psig|bara)", re.I)
_TEMP_PAT     = re.compile(r"(\d+(?:\.\d+)?)\s*(°[cf]|deg\s*[cf]|celsius|fahrenheit)", re.I)
_FLOW_PAT     = re.compile(r"(\d+(?:\.\d+)?)\s*(gpm|lpm|m3\/h|bbl\/d|scfm|nm3\/h)", re.I)

# Normalize tag separators
def _normalise_tag(raw: str) -> str:
    m = _ISA_PAT.search(raw.upper())
    if m:
        return f"{m.group(1)}-{m.group(2)}"
    return raw.upper().strip()


# ---------------------------------------------------------------------------
# Document loading
# ---------------------------------------------------------------------------

def _load_docx_paragraphs(path: Path) -> List[Tuple[str, str]]:
    """
    Load a .docx and return list of (style_name, text) tuples.
    Requires python-docx; falls back to zip/XML extraction.
    """
    try:
        import docx  # python-docx

        doc  = docx.Document(str(path))
        rows = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style else "Normal"
            rows.append((style, text))
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        rows.append(("Normal", text))
        return rows

    except ImportError:
        log.warning("python-docx not installed — falling back to zip XML extraction")
        return _load_docx_xml_fallback(path)


def _load_docx_xml_fallback(path: Path) -> List[Tuple[str, str]]:
    """Extract text from .docx using raw XML parsing (no python-docx needed)."""
    try:
        with zipfile.ZipFile(str(path), "r") as z:
            xml = z.read("word/document.xml").decode("utf-8", errors="replace")
    except Exception as e:
        log.error("Cannot open .docx zip: %s", e)
        return []

    # Strip XML tags, preserve paragraph breaks
    xml = re.sub(r"<w:p[ >]", "\n", xml)
    xml = re.sub(r"<[^>]+>", "", xml)
    # Decode XML entities
    xml = xml.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")

    rows = []
    for line in xml.split("\n"):
        text = line.strip()
        if text:
            rows.append(("Normal", text))
    return rows


def _load_text_file(path: Path) -> List[Tuple[str, str]]:
    """Load a plain .txt file."""
    rows = []
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        for line in f:
            text = line.strip()
            if text:
                style = "Heading 1" if text.isupper() and len(text) < 80 else "Normal"
                rows.append((style, text))
    return rows


# ---------------------------------------------------------------------------
# Section splitter
# ---------------------------------------------------------------------------

def _split_into_sections(
    paragraphs: List[Tuple[str, str]],
    heading_styles: List[str],
) -> List[Tuple[str, str, List[str]]]:
    """
    Group paragraphs into (heading_style, heading_text, [body_lines]).
    Returns list of section tuples.
    """
    sections: List[Tuple[str, str, List[str]]] = []
    current_heading_style = "Normal"
    current_heading = "Preamble"
    current_body: List[str] = []

    for style, text in paragraphs:
        is_heading = any(style.startswith(hs) for hs in heading_styles)
        # Also treat ALL-CAPS short lines as headings
        if not is_heading and text.isupper() and 5 < len(text) < 100:
            is_heading = True

        if is_heading:
            if current_body:
                sections.append((current_heading_style, current_heading, current_body))
            current_heading_style = style
            current_heading = text
            current_body = []
        else:
            current_body.append(text)

    if current_body:
        sections.append((current_heading_style, current_heading, current_body))

    return sections


# ---------------------------------------------------------------------------
# Per-section extraction
# ---------------------------------------------------------------------------

def _extract_tags(text: str) -> List[str]:
    tags = []
    for m in _ISA_PAT.finditer(text.upper()):
        tags.append(f"{m.group(1)}-{m.group(2)}")
    return list(dict.fromkeys(tags))  # preserve order, dedupe


def _extract_valve_positions(text: str) -> Dict[str, str]:
    positions: Dict[str, str] = {}
    for m in _OPEN_PAT.finditer(text):
        tag = _normalise_tag(m.group(1))
        positions[tag] = "open"
    for m in _CLOSE_PAT.finditer(text):
        tag = _normalise_tag(m.group(1))
        positions[tag] = "closed"
    return positions


def _extract_parameters(text: str) -> Dict[str, str]:
    params: Dict[str, str] = {}
    for m in _PRESSURE_PAT.finditer(text):
        params["pressure"] = f"{m.group(1)} {m.group(2).lower()}"
    for m in _TEMP_PAT.finditer(text):
        params["temperature"] = f"{m.group(1)} {m.group(2).lower()}"
    for m in _FLOW_PAT.finditer(text):
        params["flow"] = f"{m.group(1)} {m.group(2).lower()}"
    return params


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

class SopParser:
    """
    Parse an SOP document and expose structured process steps.

    Supports .docx (python-docx or zip fallback) and .txt.
    """

    def __init__(self, cfg: SopConfig | None = None):
        self.cfg = cfg or SopConfig()

    def parse(self, path: str | Path) -> List[SopStep]:
        """
        Parse the SOP file and return list of SopStep objects.
        """
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"SOP file not found: {path}")

        suffix = path.suffix.lower()
        if suffix == ".docx":
            paragraphs = _load_docx_paragraphs(path)
        elif suffix == ".txt":
            paragraphs = _load_text_file(path)
        else:
            log.warning("Unknown SOP format %s — treating as text", suffix)
            paragraphs = _load_text_file(path)

        sections = _split_into_sections(paragraphs, self.cfg.heading_styles)
        steps = self._sections_to_steps(sections)

        log.info(
            "SOP parsed: %d sections → %d steps, %d unique tags",
            len(sections),
            len(steps),
            len({t for s in steps for t in s.required_tags}),
        )
        return steps

    def parse_text(self, text: str) -> List[SopStep]:
        """Parse a raw string (e.g., from a VLM extraction)."""
        lines = [(("Heading 1" if ln.isupper() else "Normal"), ln.strip())
                 for ln in text.split("\n") if ln.strip()]
        sections = _split_into_sections(lines, self.cfg.heading_styles)
        return self._sections_to_steps(sections)

    def _sections_to_steps(
        self,
        sections: List[Tuple[str, str, List[str]]],
    ) -> List[SopStep]:
        steps: List[SopStep] = []
        for i, (style, heading, body_lines) in enumerate(sections):
            full_text = " ".join(body_lines)
            tags       = _extract_tags(heading + " " + full_text)
            positions  = _extract_valve_positions(full_text)
            parameters = _extract_parameters(full_text)

            step = SopStep(
                step_id=f"step_{i:04d}",
                heading=heading,
                text=full_text,
                required_tags=tags,
                valve_positions=positions,
                parameters=parameters,
            )
            steps.append(step)
        return steps

    def all_required_tags(self, steps: List[SopStep]) -> List[str]:
        """Deduplicated list of all ISA tags mentioned in the SOP."""
        seen = set()
        tags = []
        for s in steps:
            for t in s.required_tags:
                if t not in seen:
                    seen.add(t)
                    tags.append(t)
        return tags

    def all_valve_positions(self, steps: List[SopStep]) -> Dict[str, str]:
        """Merged valve-position map; later steps override earlier ones."""
        positions: Dict[str, str] = {}
        for s in steps:
            positions.update(s.valve_positions)
        return positions
